"""Local PDF and DOCX readers for electronic UPD documents."""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path

import cv2
import fitz
import numpy as np
import pytesseract
from docx import Document


@dataclass
class StructuredSourceContent:
    """Normalized text and table data read from one source document."""

    text: str
    tables: list[list[list[str]]] = field(default_factory=list)
    page_count: int = 1
    used_ocr: bool = False
    warnings: list[str] = field(default_factory=list)


def _clean_cell(value: object) -> str:
    """Normalize one table cell into a compact string."""
    if value is None:
        return ""
    return " ".join(str(value).replace("\u00a0", " ").split())


def _pixmap_to_bgr(pixmap: fitz.Pixmap) -> np.ndarray:
    """Convert a PyMuPDF pixmap into an OpenCV BGR image."""
    channels = pixmap.n
    array = np.frombuffer(pixmap.samples, dtype=np.uint8).reshape(
        pixmap.height,
        pixmap.width,
        channels,
    )
    if channels == 4:
        return cv2.cvtColor(array, cv2.COLOR_RGBA2BGR)
    if channels == 3:
        return cv2.cvtColor(array, cv2.COLOR_RGB2BGR)
    return cv2.cvtColor(array, cv2.COLOR_GRAY2BGR)


def _ocr_pdf_page(page: fitz.Page, lang: str) -> str:
    """Render and OCR one PDF page when no useful text layer exists."""
    pixmap = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0), alpha=False)
    image = _pixmap_to_bgr(pixmap)
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    normalized = cv2.normalize(gray, None, 0, 255, cv2.NORM_MINMAX)
    try:
        return pytesseract.image_to_string(
            normalized,
            lang=lang,
            config="--psm 6",
            timeout=20,
        )
    except RuntimeError:
        return ""


def _extract_pdf_tables(page: fitz.Page) -> list[list[list[str]]]:
    """Extract page tables when the installed PyMuPDF supports table finding."""
    if not hasattr(page, "find_tables"):
        return []
    try:
        finder = page.find_tables()
    except (AttributeError, ValueError, RuntimeError):
        return []

    tables: list[list[list[str]]] = []
    for table in getattr(finder, "tables", []):
        try:
            rows = table.extract()
        except (AttributeError, ValueError, RuntimeError):
            continue
        normalized = [[_clean_cell(cell) for cell in row] for row in rows]
        if normalized:
            tables.append(normalized)
    return tables


def read_pdf(
    source_path: Path,
    lang: str,
    deep_ocr: bool,
    debug_dir: Path | None = None,
) -> StructuredSourceContent:
    """Read PDF text and tables, using OCR only for image-only pages."""
    text_chunks: list[str] = []
    tables: list[list[list[str]]] = []
    used_ocr = False
    warnings: list[str] = []

    with fitz.open(source_path) as document:
        for page_index, page in enumerate(document):
            native_text = page.get_text("text").strip()
            page_text = native_text
            if len(native_text) < 20 or deep_ocr:
                ocr_text = _ocr_pdf_page(page, lang=lang)
                if ocr_text.strip():
                    used_ocr = True
                    if native_text:
                        page_text = f"{native_text}\n{ocr_text}"
                    else:
                        page_text = ocr_text
                elif len(native_text) < 20:
                    warnings.append(
                        f"Page {page_index + 1} has no useful text layer and OCR returned no text"
                    )
            text_chunks.append(page_text)
            tables.extend(_extract_pdf_tables(page))

        page_count = document.page_count

    combined = "\n".join(chunk for chunk in text_chunks if chunk).strip()
    if debug_dir:
        debug_dir.mkdir(parents=True, exist_ok=True)
        (debug_dir / "extracted_text.txt").write_text(combined, encoding="utf-8")

    return StructuredSourceContent(
        text=combined,
        tables=tables,
        page_count=page_count,
        used_ocr=used_ocr,
        warnings=warnings,
    )


def read_docx(
    source_path: Path,
    debug_dir: Path | None = None,
) -> StructuredSourceContent:
    """Read paragraphs and structured tables from one DOCX document."""
    document = Document(source_path)
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
    tables: list[list[list[str]]] = []
    table_text_chunks: list[str] = []

    for table in document.tables:
        rows = [
            [_clean_cell(cell.text) for cell in row.cells]
            for row in table.rows
        ]
        if rows:
            tables.append(rows)
            table_text_chunks.extend(" | ".join(row) for row in rows)

    combined = "\n".join(
        chunk
        for chunk in (*paragraphs, *table_text_chunks)
        if chunk.strip()
    ).strip()
    if debug_dir:
        debug_dir.mkdir(parents=True, exist_ok=True)
        (debug_dir / "extracted_text.txt").write_text(combined, encoding="utf-8")

    return StructuredSourceContent(
        text=combined,
        tables=tables,
        page_count=1,
    )
