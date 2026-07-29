"""Editable DOCX output built from anonymized text and OCR results."""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from statistics import median
from typing import Sequence

import fitz
from docx import Document
from docx.document import Document as DocumentType
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt
from PIL import Image, ImageSequence

from .config import (
    AnonymizationConfig,
    find_heading_token_range,
    mask_after_heading,
)
from .image import OcrPage, OcrWord, _choose_ocr_page
from .models import DetectedEntity, TextEntityAnalyzer, UnitProgressCallback
from .text import merge_entities


_PLAIN_LAYOUT = "plain"
_PRESERVE_LAYOUT = "preserve"
_A4_LONG_SIDE_POINTS = 841.89
_MIN_PAGE_SIDE_POINTS = 288.0
_MAX_PAGE_SIDE_POINTS = 1584.0


@dataclass(frozen=True)
class _LayoutLine:
    """One OCR line ordered by upright page coordinates."""

    words: tuple[OcrWord, ...]
    left: int
    top: int
    right: int
    bottom: int


def _mask_entities(text: str, entities: Sequence[DetectedEntity]) -> str:
    """Mask explicit entity spans while preserving whitespace."""
    characters = list(text)
    for entity in merge_entities(entities, len(text)):
        for index in range(entity.start, entity.end):
            if not characters[index].isspace():
                characters[index] = "█"
    return "".join(characters)


def _analyze_ocr_entities(
    text: str,
    analyzer: TextEntityAnalyzer,
) -> list[DetectedEntity]:
    """Run the OCR-aware analyzer and return merged spans."""
    analyze_ocr = getattr(analyzer, "analyze_ocr", None)
    raw_entities = (
        analyze_ocr(text) if callable(analyze_ocr) else analyzer.analyze(text)
    )
    return merge_entities(raw_entities, len(text))



def _append_plain_page(document: DocumentType, text: str, page_index: int) -> None:
    """Append one editable OCR page with a clear page boundary."""
    if page_index > 1:
        document.add_page_break()
    for line in text.splitlines() or [text]:
        document.add_paragraph(line)


def _save_document(document: DocumentType, destination: Path) -> None:
    """Save a DOCX with empty core metadata."""
    properties = document.core_properties
    properties.author = ""
    properties.last_modified_by = ""
    properties.title = ""
    properties.subject = ""
    properties.keywords = ""
    properties.comments = ""
    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)


def _layout_value(word: OcrWord, attribute: str, fallback: str) -> int:
    """Read an upright-layout coordinate with backward-compatible fallback."""
    value = getattr(word, attribute)
    if value is None:
        return int(getattr(word, fallback))
    return int(value)


def _word_left(word: OcrWord) -> int:
    return _layout_value(word, "layout_left", "left")


def _word_top(word: OcrWord) -> int:
    return _layout_value(word, "layout_top", "top")


def _word_width(word: OcrWord) -> int:
    return _layout_value(word, "layout_width", "width")


def _word_height(word: OcrWord) -> int:
    return _layout_value(word, "layout_height", "height")


def _make_layout_line(words: Sequence[OcrWord]) -> _LayoutLine:
    """Build one normalized line from positioned OCR words."""
    ordered = tuple(sorted(words, key=lambda word: (_word_left(word), word.start)))
    return _LayoutLine(
        words=ordered,
        left=min(_word_left(word) for word in ordered),
        top=min(_word_top(word) for word in ordered),
        right=max(_word_left(word) + _word_width(word) for word in ordered),
        bottom=max(_word_top(word) + _word_height(word) for word in ordered),
    )


def _group_layout_lines(page: OcrPage) -> list[_LayoutLine]:
    """Group OCR words into lines using Tesseract structure or geometry."""
    words = [word for word in page.words if word.text.strip()]
    if not words:
        return []

    if any(word.line_number > 0 for word in words):
        grouped: dict[tuple[int, int, int], list[OcrWord]] = {}
        for word in words:
            key = (
                word.block_number,
                word.paragraph_number,
                word.line_number,
            )
            grouped.setdefault(key, []).append(word)
        return sorted(
            (_make_layout_line(group) for group in grouped.values()),
            key=lambda line: (line.top, line.left),
        )

    lines: list[list[OcrWord]] = []
    for word in sorted(words, key=lambda item: (_word_top(item), _word_left(item))):
        center = _word_top(word) + _word_height(word) / 2
        selected: list[OcrWord] | None = None
        best_distance: float | None = None
        for candidate in lines:
            candidate_centers = [
                _word_top(item) + _word_height(item) / 2 for item in candidate
            ]
            candidate_center = median(candidate_centers)
            tolerance = max(
                5.0,
                median(_word_height(item) for item in candidate) * 0.65,
                _word_height(word) * 0.65,
            )
            distance = abs(center - candidate_center)
            if distance <= tolerance and (
                best_distance is None or distance < best_distance
            ):
                selected = candidate
                best_distance = distance
        if selected is None:
            lines.append([word])
        else:
            selected.append(word)
    return sorted(
        (_make_layout_line(group) for group in lines),
        key=lambda line: (line.top, line.left),
    )


def _mask_word(
    word: OcrWord,
    entities: Sequence[DetectedEntity],
    redact_entire_word: bool,
) -> str:
    """Mask the sensitive characters of one OCR word."""
    if redact_entire_word:
        return "".join("█" if not character.isspace() else character for character in word.text)
    characters = list(word.text)
    for entity in entities:
        start = max(word.start, entity.start)
        end = min(word.end, entity.end)
        if start >= end:
            continue
        for absolute_index in range(start, end):
            relative_index = absolute_index - word.start
            if 0 <= relative_index < len(characters) and not characters[relative_index].isspace():
                characters[relative_index] = "█"
    return "".join(characters)


def _page_dimensions_for_ocr(
    page: OcrPage,
    source_width_points: float,
    source_height_points: float,
) -> tuple[float, float]:
    """Return physical page dimensions matching the selected upright OCR view."""
    if page.rotation_degrees in {90, 270}:
        return source_height_points, source_width_points
    return source_width_points, source_height_points


def _image_dimensions_points(
    image: Image.Image,
    page: OcrPage,
    source_dpi: tuple[float, float] | None,
) -> tuple[float, float]:
    """Estimate an editable page size from image DPI or A4-scale geometry."""
    if source_dpi is not None:
        dpi_x, dpi_y = source_dpi
        if 50 <= dpi_x <= 1200 and 50 <= dpi_y <= 1200:
            width = image.width * 72.0 / dpi_x
            height = image.height * 72.0 / dpi_y
            width = min(_MAX_PAGE_SIDE_POINTS, max(_MIN_PAGE_SIDE_POINTS, width))
            height = min(_MAX_PAGE_SIDE_POINTS, max(_MIN_PAGE_SIDE_POINTS, height))
            return _page_dimensions_for_ocr(page, width, height)

    layout_width = page.layout_width or (
        image.height if page.rotation_degrees in {90, 270} else image.width
    )
    layout_height = page.layout_height or (
        image.width if page.rotation_degrees in {90, 270} else image.height
    )
    ratio = max(0.2, min(5.0, layout_width / max(layout_height, 1)))
    if ratio >= 1:
        width = _A4_LONG_SIDE_POINTS
        height = max(_MIN_PAGE_SIDE_POINTS, width / ratio)
    else:
        height = _A4_LONG_SIDE_POINTS
        width = max(_MIN_PAGE_SIDE_POINTS, height * ratio)
    return width, height


def _configure_page_section(
    document: DocumentType,
    page_index: int,
    width_points: float,
    height_points: float,
):
    """Create and configure one DOCX section for an OCR page."""
    section = (
        document.sections[0]
        if page_index == 1
        else document.add_section(WD_SECTION.NEW_PAGE)
    )
    section.page_width = Pt(width_points)
    section.page_height = Pt(height_points)
    margin = min(24.0, max(10.0, min(width_points, height_points) * 0.025))
    section.left_margin = Pt(margin)
    section.right_margin = Pt(margin)
    section.top_margin = Pt(margin)
    section.bottom_margin = Pt(margin)
    section.header_distance = Pt(0)
    section.footer_distance = Pt(0)
    return section, margin


def _append_redacted_layout_page(
    document: DocumentType,
    page_index: int,
    width_points: float,
    height_points: float,
) -> None:
    """Append a privacy-safe editable placeholder for a fully redacted page."""
    _configure_page_section(document, page_index, width_points, height_points)
    line_count = max(8, min(28, int(height_points / 32)))
    block_count = max(20, min(90, int(width_points / 7)))
    for _ in range(line_count):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        paragraph.paragraph_format.line_spacing = Pt(24)
        run = paragraph.add_run("█" * block_count)
        run.font.name = "Arial"
        run.font.size = Pt(18)


def _append_preserved_ocr_page(
    document: DocumentType,
    page: OcrPage,
    entities: Sequence[DetectedEntity],
    config: AnonymizationConfig,
    page_index: int,
    width_points: float,
    height_points: float,
) -> bool:
    """Append OCR text using approximate source coordinates and typography."""
    _section, margin = _configure_page_section(
        document,
        page_index,
        width_points,
        height_points,
    )
    lines = _group_layout_lines(page)
    if not lines:
        document.add_paragraph("")
        return False

    layout_width = page.layout_width or page.original_width or 1
    layout_height = page.layout_height or page.original_height or 1
    usable_width = max(72.0, width_points - 2 * margin)
    usable_height = max(72.0, height_points - 2 * margin)
    scale_x = usable_width / layout_width
    scale_y = usable_height / layout_height

    heading_range = find_heading_token_range(
        [word.text for word in page.words],
        config.included_paragraphs,
    )
    heading_end = heading_range[1] if heading_range is not None else None
    word_positions = {id(word): index for index, word in enumerate(page.words)}

    previous_bottom = 0
    for line in lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Pt(max(0.0, line.left * scale_x))
        paragraph.paragraph_format.right_indent = Pt(
            max(0.0, (layout_width - line.right) * scale_x)
        )
        gap_points = max(0.0, (line.top - previous_bottom) * scale_y)
        paragraph.paragraph_format.space_before = Pt(min(72.0, gap_points))
        paragraph.paragraph_format.space_after = Pt(0)

        line_heights = [_word_height(word) for word in line.words]
        line_height_points = max(7.0, median(line_heights) * scale_y)
        font_size = max(6.0, min(28.0, line_height_points * 0.82))
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        paragraph.paragraph_format.line_spacing = Pt(
            max(font_size * 1.15, line_height_points)
        )

        character_widths = [
            _word_width(word) / max(len(word.text), 1) for word in line.words
        ]
        average_character_width = max(2.0, median(character_widths))
        previous_right = line.left
        for word_index, word in enumerate(line.words):
            if word_index > 0:
                gap = max(0, _word_left(word) - previous_right)
                spaces = max(1, min(40, round(gap / average_character_width)))
                paragraph.add_run(" " * spaces)
            absolute_index = word_positions.get(id(word), -1)
            redact_entire_word = (
                heading_end is not None and absolute_index >= heading_end
            )
            run = paragraph.add_run(
                _mask_word(word, entities, redact_entire_word)
            )
            run.font.name = "Arial"
            run.font.size = Pt(font_size)
            previous_right = _word_left(word) + _word_width(word)
        previous_bottom = max(previous_bottom, line.bottom)

    return heading_range is not None


def _normalize_output_layout(output_layout: str | None) -> str:
    """Validate and normalize editable DOCX layout selection."""
    normalized = (output_layout or _PLAIN_LAYOUT).strip().lower()
    if normalized not in {_PLAIN_LAYOUT, _PRESERVE_LAYOUT}:
        raise ValueError(f"Unsupported output layout: {output_layout}")
    return normalized


def anonymize_pdf_to_docx(
    source: Path,
    destination: Path,
    analyzer: TextEntityAnalyzer,
    lang: str,
    config: AnonymizationConfig,
    progress_callback: UnitProgressCallback | None = None,
    dpi: int = 220,
    output_layout: str | None = None,
) -> int:
    """OCR a PDF and write anonymized editable text to DOCX."""
    layout = _normalize_output_layout(output_layout)
    output = Document()
    detected = 0
    redact_following_pages = False
    with fitz.open(source) as pdf:
        if pdf.needs_pass:
            raise ValueError("Password-protected PDF files are not supported")
        matrix = fitz.Matrix(dpi / 72.0, dpi / 72.0)
        for page_index, page in enumerate(pdf, start=1):
            if progress_callback is not None:
                progress_callback("page", page_index, pdf.page_count)
            source_width_points = float(page.rect.width)
            source_height_points = float(page.rect.height)
            if redact_following_pages:
                if layout == _PRESERVE_LAYOUT:
                    _append_redacted_layout_page(
                        output,
                        page_index,
                        source_width_points,
                        source_height_points,
                    )
                else:
                    _append_plain_page(output, "████████", page_index)
                continue
            pixmap = page.get_pixmap(matrix=matrix, alpha=False)
            mode = "RGBA" if pixmap.alpha else "RGB"
            image = Image.frombytes(
                mode,
                (pixmap.width, pixmap.height),
                pixmap.samples,
            ).convert("RGB")
            ocr_page, _orientation_entities = _choose_ocr_page(
                image,
                analyzer=analyzer,
                lang=lang,
                config=config,
            )
            entities = _analyze_ocr_entities(ocr_page.text, analyzer)
            detected += len(entities)
            if layout == _PRESERVE_LAYOUT:
                width_points, height_points = _page_dimensions_for_ocr(
                    ocr_page,
                    source_width_points,
                    source_height_points,
                )
                heading_found = _append_preserved_ocr_page(
                    output,
                    ocr_page,
                    entities,
                    config,
                    page_index,
                    width_points,
                    height_points,
                )
            else:
                masked = _mask_entities(ocr_page.text, entities)
                masked, heading_found = mask_after_heading(
                    ocr_page.text,
                    masked,
                    config.included_paragraphs,
                )
                _append_plain_page(output, masked, page_index)
            detected += int(heading_found)
            redact_following_pages = heading_found
    _save_document(output, destination)
    return detected


def anonymize_image_to_docx(
    source: Path,
    destination: Path,
    analyzer: TextEntityAnalyzer,
    lang: str,
    config: AnonymizationConfig,
    progress_callback: UnitProgressCallback | None = None,
    output_layout: str | None = None,
) -> int:
    """OCR raster frames and write anonymized editable text to DOCX."""
    layout = _normalize_output_layout(output_layout)
    output = Document()
    detected = 0
    redact_following_pages = False
    with Image.open(source) as image:
        frame_count = getattr(image, "n_frames", 1)
        raw_dpi = image.info.get("dpi")
        source_dpi: tuple[float, float] | None = None
        if isinstance(raw_dpi, tuple) and len(raw_dpi) >= 2:
            try:
                source_dpi = (float(raw_dpi[0]), float(raw_dpi[1]))
            except (TypeError, ValueError):
                source_dpi = None
        for frame_index, frame in enumerate(ImageSequence.Iterator(image), start=1):
            if progress_callback is not None:
                progress_callback("frame", frame_index, frame_count)
            frame_image = frame.copy().convert("RGB")
            if redact_following_pages:
                if layout == _PRESERVE_LAYOUT:
                    width_points, height_points = _image_dimensions_points(
                        frame_image,
                        OcrPage(
                            text="",
                            words=(),
                            rotation_degrees=0,
                            original_width=frame_image.width,
                            original_height=frame_image.height,
                            layout_width=frame_image.width,
                            layout_height=frame_image.height,
                        ),
                        source_dpi,
                    )
                    _append_redacted_layout_page(
                        output,
                        frame_index,
                        width_points,
                        height_points,
                    )
                else:
                    _append_plain_page(output, "████████", frame_index)
                continue
            ocr_page, _orientation_entities = _choose_ocr_page(
                frame_image,
                analyzer=analyzer,
                lang=lang,
                config=config,
            )
            entities = _analyze_ocr_entities(ocr_page.text, analyzer)
            detected += len(entities)
            if layout == _PRESERVE_LAYOUT:
                width_points, height_points = _image_dimensions_points(
                    frame_image,
                    ocr_page,
                    source_dpi,
                )
                heading_found = _append_preserved_ocr_page(
                    output,
                    ocr_page,
                    entities,
                    config,
                    frame_index,
                    width_points,
                    height_points,
                )
            else:
                masked = _mask_entities(ocr_page.text, entities)
                masked, heading_found = mask_after_heading(
                    ocr_page.text,
                    masked,
                    config.included_paragraphs,
                )
                _append_plain_page(output, masked, frame_index)
            detected += int(heading_found)
            redact_following_pages = heading_found
    _save_document(output, destination)
    return detected


def anonymize_text_to_docx(
    text: str,
    destination: Path,
    analyzer: TextEntityAnalyzer,
    config: AnonymizationConfig,
) -> int:
    """Write anonymized plain text into an editable DOCX."""
    entities = merge_entities(analyzer.analyze(text), len(text))
    masked = _mask_entities(text, entities)
    masked, heading_found = mask_after_heading(
        text,
        masked,
        config.included_paragraphs,
    )
    output = Document()
    _append_plain_page(output, masked, 1)
    _save_document(output, destination)
    return len(entities) + int(heading_found)
