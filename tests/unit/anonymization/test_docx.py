"""Tests for DOCX package anonymization and fail-closed behavior."""

from __future__ import annotations

import zipfile
from pathlib import Path

import pytest
from docx import Document

from source_docs_processor.anonymization.docx import anonymize_docx_file
from source_docs_processor.anonymization.models import DetectedEntity


class NameAnalyzer:
    """Detect one fictional full name wherever it occurs."""

    def analyze(self, text: str) -> list[DetectedEntity]:
        """Return the fictional name span for deterministic tests."""
        value = "Иван Петров"
        start = text.find(value)
        if start < 0:
            return []
        return [DetectedEntity(start, start + len(value), "PERSON")]


def test_docx_anonymization_masks_text_and_clears_core_properties(tmp_path: Path) -> None:
    """Verify visible DOCX text and author metadata are sanitized.

    Protected risk: anonymizing only body text would leave personal data in
    package metadata or preserve the original name in the rendered document.
    """
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    document = Document()
    document.add_paragraph("Получатель: Иван Петров")
    document.core_properties.author = "Иван Петров"
    document.save(source)

    detected = anonymize_docx_file(source, output, NameAnalyzer())

    anonymized = Document(output)
    assert detected >= 1
    assert "Иван Петров" not in anonymized.paragraphs[0].text
    assert anonymized.core_properties.author in {None, ""}


def test_docx_anonymization_rejects_embedded_binary_content(tmp_path: Path) -> None:
    """Verify unsupported embedded objects fail instead of being copied unchanged.

    Protected risk: opaque OLE or embedded workbook data may retain unredacted
    personal information inside an apparently anonymized DOCX file.
    """
    source = tmp_path / "unsafe.docx"
    output = tmp_path / "output.docx"
    with zipfile.ZipFile(source, "w") as archive:
        archive.writestr("[Content_Types].xml", "<Types xmlns='http://schemas.openxmlformats.org/package/2006/content-types'/>")
        archive.writestr("word/embeddings/object1.bin", b"private")

    with pytest.raises(ValueError, match="embedded or active content"):
        anonymize_docx_file(source, output, NameAnalyzer())

    assert not output.exists()


def test_docx_anonymization_replaces_text_across_runs(tmp_path: Path) -> None:
    """Verify configured replacement can span multiple formatted DOCX runs.

    Protected risk: Word may split one surname across runs, and replacing only
    inside individual runs would leave the original private value visible.
    """
    from source_docs_processor.anonymization.config import (
        AnonymizationConfig,
        ConfiguredTextAnalyzer,
        ReplacementRule,
    )

    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("Получатель: Васи")
    paragraph.add_run("льев")
    document.save(source)
    config = AnonymizationConfig(
        included_and_replaced=(ReplacementRule("Васильев", "Иванов"),)
    )

    detected = anonymize_docx_file(
        source,
        output,
        ConfiguredTextAnalyzer(None, config),
        config=config,
    )

    anonymized = Document(output)
    assert detected == 1
    assert anonymized.paragraphs[0].text == "Получатель: Иванов"
