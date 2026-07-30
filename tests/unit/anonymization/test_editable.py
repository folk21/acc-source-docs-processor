"""Tests for editable anonymization output."""

from __future__ import annotations

from pathlib import Path

import fitz
import pytest
from docx import Document

from source_docs_processor.features.anonymization.config import AnonymizationConfig
from source_docs_processor.features.anonymization.editable import anonymize_pdf_to_docx
from source_docs_processor.features.anonymization.image import OcrPage
from source_docs_processor.features.anonymization.models import DetectedEntity


class IncludedAnalyzer:
    """Detect one configured fictional word in OCR text."""

    def analyze(self, text: str) -> list[DetectedEntity]:
        """Return the fictional word span."""
        start = text.find("Квантовая")
        if start < 0:
            return []
        return [DetectedEntity(start, start + len("Квантовая"), "CONFIG_INCLUDED")]

    def analyze_ocr(self, text: str) -> list[DetectedEntity]:
        """Use the same deterministic result for OCR text."""
        return self.analyze(text)


def test_scanned_pdf_can_be_reconstructed_as_editable_docx(
    tmp_path: Path,
    monkeypatch,
) -> None:
    """Verify PDF-to-DOCX output contains editable masked OCR text.

    Protected risk: editable output must not embed the original scan or expose
    the configured literal in a hidden text layer.
    """
    source = tmp_path / "source.pdf"
    destination = tmp_path / "output.docx"
    pdf = fitz.open()
    pdf.new_page(width=300, height=200)
    pdf.save(source)
    pdf.close()

    fake_page = OcrPage(
        text="Раздел 2.2.1 Квантовая долина",
        words=(),
        rotation_degrees=0,
        original_width=300,
        original_height=200,
    )
    monkeypatch.setattr(
        "source_docs_processor.features.anonymization.editable._choose_ocr_page",
        lambda *args, **kwargs: (fake_page, []),
    )

    detected = anonymize_pdf_to_docx(
        source,
        destination,
        IncludedAnalyzer(),
        lang="rus+eng",
        config=AnonymizationConfig(included=("Квантовая",)),
    )

    text = "\n".join(paragraph.text for paragraph in Document(destination).paragraphs)
    assert detected == 1
    assert "Квантовая" not in text
    assert "Раздел 2.2.1" in text
    assert "долина" in text


def test_preserve_layout_reconstructs_page_geometry_and_positioned_lines(
    tmp_path: Path,
    monkeypatch,
) -> None:
    """Verify preserve mode writes positioned editable text without source images.

    Protected risk: layout reconstruction must retain useful page geometry while
    never embedding the original scan that may still contain private pixels.
    """
    from zipfile import ZipFile

    from source_docs_processor.features.anonymization.image import OcrWord

    source = tmp_path / "landscape.pdf"
    destination = tmp_path / "output.docx"
    pdf = fitz.open()
    pdf.new_page(width=420, height=280)
    pdf.save(source)
    pdf.close()

    text = "Раздел 2.2.1 Квантовая долина"
    values = ("Раздел", "2.2.1", "Квантовая", "долина")
    lefts = (30, 105, 180, 300)
    words = []
    offset = 0
    for index, (value, left) in enumerate(zip(values, lefts), start=1):
        start = text.index(value, offset)
        end = start + len(value)
        offset = end
        words.append(
            OcrWord(
                text=value,
                start=start,
                end=end,
                left=left,
                top=70,
                width=max(45, len(value) * 9),
                height=20,
                confidence=90.0,
                layout_left=left,
                layout_top=70,
                layout_width=max(45, len(value) * 9),
                layout_height=20,
                block_number=1,
                paragraph_number=1,
                line_number=1,
            )
        )
    fake_page = OcrPage(
        text=text,
        words=tuple(words),
        rotation_degrees=0,
        original_width=420,
        original_height=280,
        layout_width=420,
        layout_height=280,
    )
    monkeypatch.setattr(
        "source_docs_processor.features.anonymization.editable._choose_ocr_page",
        lambda *args, **kwargs: (fake_page, []),
    )

    detected = anonymize_pdf_to_docx(
        source,
        destination,
        IncludedAnalyzer(),
        lang="rus+eng",
        config=AnonymizationConfig(included=("Квантовая",)),
        output_layout="preserve",
    )

    output = Document(destination)
    document_text = "\n".join(paragraph.text for paragraph in output.paragraphs)
    positioned = next(paragraph for paragraph in output.paragraphs if "Раздел" in paragraph.text)
    assert detected == 1
    assert "Квантовая" not in document_text
    assert "Раздел" in document_text
    assert "2.2.1" in document_text
    assert positioned.paragraph_format.left_indent is not None
    assert positioned.paragraph_format.left_indent.pt > 10
    assert output.sections[0].page_width.pt == pytest.approx(420, abs=1)
    assert output.sections[0].page_height.pt == pytest.approx(280, abs=1)
    with ZipFile(destination) as archive:
        assert not any(name.startswith("word/media/") for name in archive.namelist())


def test_preserve_layout_swaps_page_geometry_for_rotated_ocr(
    tmp_path: Path,
    monkeypatch,
) -> None:
    """Verify upright OCR rotation also rotates the editable DOCX page.

    Protected risk: a sideways scan must not produce upright text on a page whose
    dimensions still describe the sideways source orientation.
    """
    from source_docs_processor.features.anonymization.image import OcrWord

    source = tmp_path / "rotated.pdf"
    destination = tmp_path / "output.docx"
    pdf = fitz.open()
    pdf.new_page(width=280, height=420)
    pdf.save(source)
    pdf.close()

    word = OcrWord(
        text="Текст",
        start=0,
        end=5,
        left=10,
        top=10,
        width=40,
        height=18,
        confidence=90.0,
        layout_left=20,
        layout_top=30,
        layout_width=40,
        layout_height=18,
        block_number=1,
        paragraph_number=1,
        line_number=1,
    )
    fake_page = OcrPage(
        text="Текст",
        words=(word,),
        rotation_degrees=90,
        original_width=280,
        original_height=420,
        layout_width=420,
        layout_height=280,
    )
    monkeypatch.setattr(
        "source_docs_processor.features.anonymization.editable._choose_ocr_page",
        lambda *args, **kwargs: (fake_page, []),
    )

    anonymize_pdf_to_docx(
        source,
        destination,
        IncludedAnalyzer(),
        lang="rus+eng",
        config=AnonymizationConfig(),
        output_layout="preserve",
    )

    output = Document(destination)
    assert output.sections[0].page_width.pt == pytest.approx(420, abs=1)
    assert output.sections[0].page_height.pt == pytest.approx(280, abs=1)


def test_preserve_layout_writes_replacement_text_instead_of_mask(
    tmp_path: Path,
    monkeypatch,
) -> None:
    """Verify preserve-layout DOCX contains the configured pseudonym as text.

    Protected risk: editable output must retain the replacement value rather than
    emitting opaque blocks or the original sensitive OCR token.
    """
    from source_docs_processor.features.anonymization.config import (
        AnonymizationConfig,
        ConfiguredTextAnalyzer,
        ReplacementRule,
    )
    from source_docs_processor.features.anonymization.image import OcrWord

    source = tmp_path / "source.pdf"
    destination = tmp_path / "output.docx"
    pdf = fitz.open()
    pdf.new_page(width=420, height=280)
    pdf.save(source)
    pdf.close()

    source_value = "Квантовая"
    word = OcrWord(
        text=source_value,
        start=0,
        end=len(source_value),
        left=40,
        top=70,
        width=100,
        height=20,
        confidence=90.0,
        layout_left=40,
        layout_top=70,
        layout_width=100,
        layout_height=20,
        block_number=1,
        paragraph_number=1,
        line_number=1,
    )
    fake_page = OcrPage(
        text=source_value,
        words=(word,),
        rotation_degrees=0,
        original_width=420,
        original_height=280,
        layout_width=420,
        layout_height=280,
    )
    monkeypatch.setattr(
        "source_docs_processor.features.anonymization.editable._choose_ocr_page",
        lambda *args, **kwargs: (fake_page, []),
    )
    config = AnonymizationConfig(
        included_and_replaced=(ReplacementRule(source_value, "цифровая"),)
    )

    detected = anonymize_pdf_to_docx(
        source,
        destination,
        ConfiguredTextAnalyzer(None, config),
        lang="rus+eng",
        config=config,
        output_layout="preserve",
    )

    output = Document(destination)
    document_text = "\n".join(paragraph.text for paragraph in output.paragraphs)
    assert detected == 1
    assert source_value not in document_text
    assert "цифровая" in document_text
