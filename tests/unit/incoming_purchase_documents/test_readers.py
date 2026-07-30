from pathlib import Path

import fitz
from docx import Document

from source_docs_processor.features.document_types.incoming_purchase_documents.readers import (
    read_docx,
    read_pdf,
)


def test_pdf_reader_prefers_native_text_without_ocr(tmp_path):
    """Verify text PDFs are read locally without invoking OCR fallback.

    Protected risk: clean supplier PDFs should preserve exact text and avoid the
    slower and less accurate OCR path.
    """
    pdf_path = tmp_path / "upd.pdf"
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "UPD status 1 native PDF text")
    document.save(pdf_path)
    document.close()

    content = read_pdf(pdf_path, lang="rus+eng", deep_ocr=False)

    assert "UPD status 1 native PDF text" in content.text
    assert content.used_ocr is False
    assert content.page_count == 1


def test_docx_reader_preserves_paragraphs_and_table_rows(tmp_path):
    """Verify DOCX tables remain structured for later item extraction.

    Protected risk: flattening Word tables into text would make goods columns
    unreliable and prevent deterministic row mapping.
    """
    docx_path = tmp_path / "upd.docx"
    document = Document()
    document.add_paragraph("Универсальный передаточный документ")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Наименование товара"
    table.cell(0, 1).text = "Количество"
    table.cell(1, 0).text = "Учебный товар"
    table.cell(1, 1).text = "2"
    document.save(docx_path)

    content = read_docx(docx_path)

    assert "Универсальный передаточный документ" in content.text
    assert content.tables[0][0] == ["Наименование товара", "Количество"]
    assert content.tables[0][1] == ["Учебный товар", "2"]

