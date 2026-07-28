from pathlib import Path
from zipfile import ZipFile

from docx import Document

from source_docs_processor.cli import process_folder


def _write_upd_docx(path: Path) -> None:
    """Create a synthetic electronic UPD with one goods line."""
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    for paragraph in (
        "Универсальный передаточный документ",
        "Счет-фактура № 511 от 21.03.2026",
        "Статус: 1",
        "Продавец: ООО Учебный Поставщик",
        "ИНН/КПП продавца: 7801234567 / 780101001",
        "Покупатель: ООО Учебный Покупатель",
        "ИНН/КПП покупателя: 7701234567 / 770101001",
    ):
        document.add_paragraph(paragraph)

    table = document.add_table(rows=3, cols=8)
    headers = (
        "Наименование товара (описание выполненных работ, оказанных услуг), имущественного права",
        "Единица измерения",
        "Количество (объем)",
        "Цена (тариф) за единицу измерения",
        "Стоимость товаров (работ, услуг), имущественных прав без налога - всего",
        "Налоговая ставка",
        "Сумма налога, предъявляемая покупателю",
        "Стоимость товаров (работ, услуг), имущественных прав с налогом - всего",
    )
    item = (
        "Учебный товар",
        "шт",
        "2",
        "1000,00",
        "2000,00",
        "20%",
        "400,00",
        "2400,00",
    )
    totals = (
        "Всего к оплате",
        "",
        "",
        "",
        "2000,00",
        "",
        "400,00",
        "2400,00",
    )
    for column_index, value in enumerate(headers):
        table.cell(0, column_index).text = value
    for column_index, value in enumerate(item):
        table.cell(1, column_index).text = value
    for column_index, value in enumerate(totals):
        table.cell(2, column_index).text = value
    document.save(path)


def _archive_text(path: Path, member: str) -> str:
    """Read one UTF-8 XML member from an XLSX archive."""
    with ZipFile(path) as archive:
        return archive.read(member).decode("utf-8")


def test_registered_incoming_purchase_workflow_creates_task_workbook(tmp_path):
    """Verify electronic UPD files are referenced without unnecessary copies.

    Protected risk: explicit output directories must not gain an extra nested
    folder, and the task workbook must expose a binary dropdown while keeping
    internal task identifiers hidden.
    """
    source_dir = tmp_path / "source"
    output_dir = tmp_path / "output"
    source_file = source_dir / "supplier" / "upd_511.docx"
    _write_upd_docx(source_file)

    found, all_documents = process_folder(
        source_dir=source_dir,
        output_dir=output_dir,
        lang="rus+eng",
        document_type="incoming_purchase_documents",
    )

    workbook_path = output_dir / "реестр_упд_для_ввода_в_1с.xlsx"
    report_path = output_dir / "упд_для_ввода_в_1с_report.txt"

    assert len(found) == 1
    assert len(all_documents) == 1
    assert workbook_path.exists()
    assert report_path.exists()
    assert not (output_dir / "упд_для_ввода_в_1с").exists()
    assert not (output_dir / "documents").exists()
    assert found[0].destination_path is None
    assert found[0].items[0].name == "Учебный товар"

    workbook_xml = _archive_text(workbook_path, "xl/workbook.xml")
    shared_strings = _archive_text(workbook_path, "xl/sharedStrings.xml")
    document_sheet = _archive_text(workbook_path, "xl/worksheets/sheet1.xml")
    comments = _archive_text(workbook_path, "xl/comments1.xml")

    assert "Documents" in workbook_xml
    assert "Items" in workbook_xml
    assert "Review" in workbook_xml
    assert "_metadata" in workbook_xml
    assert 'state="hidden"' in workbook_xml
    assert "обработано" in shared_strings
    assert "Нет" in shared_strings
    assert "Учебный товар" in shared_strings
    assert "2400" in document_sheet
    assert "dataValidations" in document_sheet
    assert "Нет,Да" in document_sheet
    assert 'hidden="1"' in document_sheet
    assert "Internal stable task identifier. Do not edit." in comments


def test_upd_files_workflow_does_not_overwrite_existing_workbook(tmp_path):
    """Verify repeated runs preserve an accountant's existing task workbook.

    Protected risk: rerunning extraction after checkboxes were edited must not
    destroy the prior workbook and its manually maintained processing state.
    """
    source_dir = tmp_path / "source"
    output_dir = tmp_path / "output"
    _write_upd_docx(source_dir / "upd_511.docx")

    for _ in range(2):
        process_folder(
            source_dir=source_dir,
            output_dir=output_dir,
            lang="rus+eng",
            document_type="incoming_purchase_documents",
        )

    workbooks = sorted(output_dir.glob("реестр_упд_для_ввода_в_1с*.xlsx"))

    assert len(workbooks) == 2
    assert workbooks[0].name == "реестр_упд_для_ввода_в_1с.xlsx"
    assert workbooks[1].name == "реестр_упд_для_ввода_в_1с_2.xlsx"


def test_incoming_purchase_workflow_allows_output_in_source_directory(tmp_path):
    """Verify direct output in the source folder does not exclude input files.

    Protected risk: removing the mandatory nested output directory must not make
    every source file disappear when users choose the source folder as output.
    """
    source_dir = tmp_path / "source"
    _write_upd_docx(source_dir / "upd_511.docx")

    found, all_documents = process_folder(
        source_dir=source_dir,
        output_dir=source_dir,
        lang="rus+eng",
        document_type="incoming_purchase_documents",
    )

    assert len(found) == 1
    assert len(all_documents) == 1
    assert (source_dir / "реестр_упд_для_ввода_в_1с.xlsx").exists()
