"""Integration tests for recursive folder anonymization."""

from __future__ import annotations

from pathlib import Path

from source_docs_processor.anonymization.models import DetectedEntity
from source_docs_processor.anonymization.workflow import anonymize_folder


class FictionalNameAnalyzer:
    """Detect a fictional name in plain text fixtures."""

    def analyze(self, text: str) -> list[DetectedEntity]:
        """Return every fictional full-name occurrence."""
        value = "Иван Петров"
        entities: list[DetectedEntity] = []
        offset = 0
        while True:
            start = text.find(value, offset)
            if start < 0:
                return entities
            entities.append(
                DetectedEntity(start, start + len(value), "PERSON")
            )
            offset = start + len(value)


def test_folder_anonymization_preserves_relative_names_and_reports_unsupported_files(
    tmp_path: Path,
) -> None:
    """Verify directory-to-directory output mirrors source names and fails closed.

    Protected risk: unsupported files must not be copied unchanged into a folder
    which users may trust as fully anonymized.
    """
    source = tmp_path / "source"
    output = tmp_path / "output"
    text_path = source / "nested" / "note.txt"
    text_path.parent.mkdir(parents=True)
    text_path.write_text("Контакт: Иван Петров", encoding="utf-8")
    (source / "raw.bin").write_bytes(b"Ivan Petrov")

    summary = anonymize_folder(source, output, FictionalNameAnalyzer())

    anonymized_text = output / "nested" / "note.txt"
    assert anonymized_text.exists()
    assert "Иван Петров" not in anonymized_text.read_text(encoding="utf-8")
    assert not (output / "raw.bin").exists()
    assert summary.succeeded_count == 1
    assert summary.failed_count == 1


def test_folder_anonymization_emits_file_progress_events(tmp_path: Path) -> None:
    """Verify callers can display progress before a long file finishes.

    Protected risk: buffering all status output until the complete folder ends
    makes long OCR runs appear stalled.
    """
    source = tmp_path / "source"
    output = tmp_path / "output"
    source.mkdir()
    (source / "note.txt").write_text("Контакт: Иван Петров", encoding="utf-8")
    events = []

    anonymize_folder(
        source,
        output,
        FictionalNameAnalyzer(),
        progress_callback=events.append,
    )

    assert [event.event for event in events] == ["file_started", "file_finished"]
    assert events[0].file_index == 1
    assert events[0].file_count == 1
    assert events[-1].error is None


def test_folder_anonymization_can_write_editable_docx(tmp_path: Path) -> None:
    """Verify explicit DOCX output converts supported text files to editable documents.

    Protected risk: requesting editable output must change the extension and keep
    anonymized text editable instead of silently producing a raster PDF.
    """
    from docx import Document

    source = tmp_path / "source"
    output = tmp_path / "output"
    text_path = source / "nested" / "note.txt"
    text_path.parent.mkdir(parents=True)
    text_path.write_text("Контакт: Иван Петров", encoding="utf-8")

    summary = anonymize_folder(
        source,
        output,
        FictionalNameAnalyzer(),
        output_document_type="docx",
    )

    destination = output / "nested" / "note.docx"
    assert summary.failed_count == 0
    assert destination.exists()
    document_text = "\n".join(
        paragraph.text for paragraph in Document(destination).paragraphs
    )
    assert "Иван Петров" not in document_text
    assert "Контакт:" in document_text


def test_docx_output_avoids_same_stem_collisions(tmp_path: Path) -> None:
    """Verify converted files with the same stem receive deterministic names.

    Protected risk: a PDF and TXT with the same stem must not overwrite each
    other when both are converted to DOCX in one output folder.
    """
    source = tmp_path / "source"
    output = tmp_path / "output"
    source.mkdir()
    (source / "note.txt").write_text("Контакт: Иван Петров", encoding="utf-8")
    (source / "note.pdf").write_bytes(b"not a real PDF")

    summary = anonymize_folder(
        source,
        output,
        FictionalNameAnalyzer(),
        output_document_type="docx",
    )

    assert (output / "note__txt.docx").exists()
    assert not (output / "note.docx").exists()
    assert summary.succeeded_count == 1
    assert summary.failed_count == 1


def test_preserve_layout_requires_docx_output(tmp_path: Path) -> None:
    """Verify layout reconstruction cannot be selected for source-format output.

    Protected risk: silently ignoring preserve mode would make the CLI appear to
    produce editable layout-preserving documents when it did not.
    """
    import pytest

    source = tmp_path / "source"
    output = tmp_path / "output"
    source.mkdir()
    (source / "note.txt").write_text("Контакт: Иван Петров", encoding="utf-8")

    with pytest.raises(ValueError, match="requires --outputDocumentType docx"):
        anonymize_folder(
            source,
            output,
            FictionalNameAnalyzer(),
            output_layout="preserve",
        )


def test_dual_output_writes_source_format_and_editable_docx(tmp_path: Path) -> None:
    """Verify one source can produce anonymized source-format and DOCX outputs.

    Protected risk: requesting an editable copy must not remove the safer
    source-format anonymized artifact needed for visual review.
    """
    from docx import Document

    source = tmp_path / "source"
    output = tmp_path / "output"
    text_path = source / "nested" / "note.txt"
    text_path.parent.mkdir(parents=True)
    text_path.write_text("Контакт: Иван Петров", encoding="utf-8")

    events = []
    summary = anonymize_folder(
        source,
        output,
        FictionalNameAnalyzer(),
        output_document_type="docx",
        output_layout="preserve",
        also_output_source_format=True,
        progress_callback=events.append,
    )

    source_format = output / "nested" / "note.txt"
    editable = output / "nested" / "note.docx"
    assert summary.failed_count == 0
    assert summary.succeeded_count == 1
    assert summary.generated_files_count == 2
    assert source_format.exists()
    assert editable.exists()
    assert "Иван Петров" not in source_format.read_text(encoding="utf-8")
    editable_text = "\n".join(
        paragraph.text for paragraph in Document(editable).paragraphs
    )
    assert "Иван Петров" not in editable_text
    assert set(summary.results[0].output_paths) == {source_format, editable}
    assert events[-1].output_count == 2


def test_dual_output_does_not_duplicate_matching_docx_format(tmp_path: Path) -> None:
    """Verify DOCX input creates one artifact when DOCX is also requested.

    Protected risk: dual-output mode must not create two identical DOCX files
    or invent a confusing duplicate name for an already matching source type.
    """
    from docx import Document

    source = tmp_path / "source"
    output = tmp_path / "output"
    source.mkdir()
    document = Document()
    document.add_paragraph("Контакт: Иван Петров")
    document.save(source / "note.docx")

    summary = anonymize_folder(
        source,
        output,
        FictionalNameAnalyzer(),
        output_document_type="docx",
        also_output_source_format=True,
    )

    assert summary.failed_count == 0
    assert summary.generated_files_count == 1
    assert (output / "note.docx").exists()
    assert list(output.glob("*.docx")) == [output / "note.docx"]


def test_dual_output_requires_explicit_target_type(tmp_path: Path) -> None:
    """Verify source-format duplication cannot be enabled without conversion.

    Protected risk: accepting the flag alone would create ambiguous duplicate
    outputs without a second requested document format.
    """
    import pytest

    source = tmp_path / "source"
    output = tmp_path / "output"
    source.mkdir()
    (source / "note.txt").write_text("Контакт", encoding="utf-8")

    with pytest.raises(ValueError, match="requires --outputDocumentType"):
        anonymize_folder(
            source,
            output,
            FictionalNameAnalyzer(),
            also_output_source_format=True,
        )


def test_dual_output_avoids_collision_with_existing_docx_source(
    tmp_path: Path,
    monkeypatch,
) -> None:
    """Verify converted DOCX names do not overwrite source-format DOCX output.

    Protected risk: a PDF and DOCX sharing one stem must both survive dual-output
    mode without the converted PDF replacing the anonymized DOCX source file.
    """
    from docx import Document
    from source_docs_processor.anonymization import workflow

    source = tmp_path / "source"
    output = tmp_path / "output"
    source.mkdir()
    document = Document()
    document.add_paragraph("Контакт: Иван Петров")
    document.save(source / "note.docx")
    (source / "note.pdf").write_bytes(b"synthetic PDF placeholder")

    def fake_atomic_anonymize(source_path, destination, **_kwargs):
        destination.parent.mkdir(parents=True, exist_ok=True)
        destination.write_bytes(source_path.read_bytes())
        return 1

    monkeypatch.setattr(workflow, "_atomic_anonymize", fake_atomic_anonymize)

    summary = anonymize_folder(
        source,
        output,
        FictionalNameAnalyzer(),
        output_document_type="docx",
        also_output_source_format=True,
    )

    assert (output / "note.docx").exists()
    assert (output / "note.pdf").exists()
    assert (output / "note__pdf.docx").exists()
    assert summary.succeeded_count == 2
    assert summary.failed_count == 0
    assert summary.generated_files_count == 3


def test_dual_output_removes_both_variants_when_conversion_fails(
    tmp_path: Path,
    monkeypatch,
) -> None:
    """Verify a failed second variant removes the completed first variant.

    Protected risk: a partially successful dual-output run must not leave a
    source-format artifact that users could mistake for a complete result set.
    """
    from source_docs_processor.anonymization import workflow

    source = tmp_path / "source"
    output = tmp_path / "output"
    source.mkdir()
    (source / "note.txt").write_text("Контакт: Иван Петров", encoding="utf-8")

    original_atomic = workflow._atomic_anonymize

    def fail_docx_variant(source_path, destination, **kwargs):
        if kwargs.get("output_document_type") == "docx":
            raise RuntimeError("synthetic DOCX conversion failure")
        return original_atomic(source_path, destination, **kwargs)

    monkeypatch.setattr(workflow, "_atomic_anonymize", fail_docx_variant)

    summary = anonymize_folder(
        source,
        output,
        FictionalNameAnalyzer(),
        output_document_type="docx",
        also_output_source_format=True,
    )

    assert summary.succeeded_count == 0
    assert summary.failed_count == 1
    assert summary.generated_files_count == 0
    assert not (output / "note.txt").exists()
    assert not (output / "note.docx").exists()


def test_folder_anonymization_replaces_configured_text_in_both_outputs(
    tmp_path: Path,
) -> None:
    """Verify replacement rules are applied to source-format and DOCX artifacts.

    Protected risk: dual output must not apply pseudonyms to only one variant and
    leave the original sensitive literal in the other artifact.
    """
    from docx import Document

    from source_docs_processor.anonymization.config import (
        AnonymizationConfig,
        ConfiguredTextAnalyzer,
        ReplacementRule,
    )

    source = tmp_path / "source"
    output = tmp_path / "output"
    source.mkdir()
    (source / "note.txt").write_text(
        "Ответственный: Васильев",
        encoding="utf-8",
    )
    config = AnonymizationConfig(
        included_and_replaced=(ReplacementRule("Васильев", "Иванов"),)
    )

    summary = anonymize_folder(
        source,
        output,
        ConfiguredTextAnalyzer(None, config),
        config=config,
        output_document_type="docx",
        also_output_source_format=True,
    )

    source_text = (output / "note.txt").read_text(encoding="utf-8")
    editable_text = "\n".join(
        paragraph.text
        for paragraph in Document(output / "note.docx").paragraphs
    )
    assert summary.failed_count == 0
    assert "Васильев" not in source_text
    assert "Васильев" not in editable_text
    assert "Иванов" in source_text
    assert "Иванов" in editable_text

def test_output_ancestor_does_not_exclude_nested_source_files(
    tmp_path: Path,
    monkeypatch,
) -> None:
    """Verify a current output directory may contain the source directory.

    Protected risk: treating every file below the output root as generated output
    previously excluded the complete nested source tree and produced zero files
    without reporting an error.
    """
    workspace = tmp_path / "workspace"
    source = workspace / "input"
    source.mkdir(parents=True)
    (source / "note.txt").write_text("Контакт: Иван Петров", encoding="utf-8")
    monkeypatch.chdir(workspace)

    summary = anonymize_folder(
        Path("input"),
        Path("."),
        FictionalNameAnalyzer(),
    )

    destination = workspace / "note.txt"
    assert summary.succeeded_count == 1
    assert summary.failed_count == 0
    assert destination.exists()
    assert "Иван Петров" not in destination.read_text(encoding="utf-8")


def test_empty_source_directory_fails_with_resolved_path_diagnostics(
    tmp_path: Path,
) -> None:
    """Verify an empty scan cannot finish successfully with zero generated files.

    Protected risk: a path or exclusion mistake must produce a clear non-zero
    failure instead of a misleading successful summary with an empty output.
    """
    import pytest

    source = tmp_path / "source"
    output = tmp_path / "output"
    source.mkdir()

    with pytest.raises(ValueError, match="No source files were found") as error:
        anonymize_folder(source, output, FictionalNameAnalyzer())

    message = str(error.value)
    assert str(source.resolve()) in message
    assert str(output.resolve()) in message



def test_nested_output_directory_remains_excluded_from_source_scan(
    tmp_path: Path,
) -> None:
    """Verify generated output below the source tree is not reprocessed.

    Protected risk: fixing ancestor-output traversal must not remove the guard
    that prevents files already inside a nested output directory from becoming
    new anonymization inputs.
    """
    source = tmp_path / "source"
    output = source / "output"
    source.mkdir()
    output.mkdir()
    (source / "note.txt").write_text("Контакт: Иван Петров", encoding="utf-8")
    existing_output = output / "previous.txt"
    existing_output.write_text("previous result", encoding="utf-8")

    summary = anonymize_folder(source, output, FictionalNameAnalyzer())

    assert summary.succeeded_count == 1
    assert summary.failed_count == 0
    assert [result.source_path for result in summary.results] == [source / "note.txt"]
    assert (output / "note.txt").exists()
    assert existing_output.read_text(encoding="utf-8") == "previous result"


def test_clear_output_preserves_directory_inode_and_removes_old_files(
    tmp_path: Path,
) -> None:
    """Verify output cleanup keeps an open terminal attached to the live directory.

    Protected risk: deleting and recreating the output root leaves another macOS
    terminal inside an unlinked directory that cannot see newly generated files.
    """
    import os

    source = tmp_path / "source"
    output = tmp_path / "output"
    source.mkdir()
    output.mkdir()
    (source / "note.txt").write_text("Контакт: Иван Петров", encoding="utf-8")
    (output / "old.txt").write_text("old result", encoding="utf-8")
    nested = output / "nested"
    nested.mkdir()
    (nested / "old.txt").write_text("old nested result", encoding="utf-8")

    output_descriptor = os.open(output, os.O_RDONLY)
    nested_inode = nested.stat().st_ino
    try:
        original_inode = os.fstat(output_descriptor).st_ino
        summary = anonymize_folder(
            source,
            output,
            FictionalNameAnalyzer(),
            clear_output=True,
        )

        assert summary.failed_count == 0
        assert os.fstat(output_descriptor).st_ino == original_inode
        assert output.stat().st_ino == original_inode
        assert nested.stat().st_ino == nested_inode
        assert not (output / "old.txt").exists()
        assert not (nested / "old.txt").exists()
        assert (output / "note.txt").exists()
    finally:
        os.close(output_descriptor)


def test_clear_output_rejects_source_inside_output(tmp_path: Path) -> None:
    """Verify cleanup cannot erase a source tree nested below the output root.

    Protected risk: a convenience cleanup flag must fail before deleting input
    documents when output is an ancestor of source.
    """
    import pytest

    output = tmp_path / "output"
    source = output / "source"
    source.mkdir(parents=True)
    (source / "note.txt").write_text("Контакт: Иван Петров", encoding="utf-8")

    with pytest.raises(ValueError, match="source directory is inside"):
        anonymize_folder(
            source,
            output,
            FictionalNameAnalyzer(),
            clear_output=True,
        )

    assert (source / "note.txt").exists()
